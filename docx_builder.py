from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        self.set_default_font()

    def set_default_font(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

    def is_arabic(self, text):
        # A simple check for Arabic characters
        return bool(re.search(r'[\u0600-\u06FF]', text))

    def add_heading(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        if self.is_arabic(text):
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_rtl(heading)
        return heading

    def set_rtl(self, obj):
        """Sets Right-To-Left direction for a paragraph or run."""
        from docx.oxml import OxmlElement
        
        # For the XML level RTL support
        p = obj._element
        pPr = p.get_or_add_pPr()
        
        # Add bidi element
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        
        # Also set text alignment to right
        if hasattr(obj, 'paragraph_format'):
            obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_paragraph(self, text):
        if not text:
            return
        
        p = self.doc.add_paragraph()
        if self.is_arabic(text):
            self.set_rtl(p)
        
        p.add_run(text)
        return p

    def build_from_markdown(self, markdown_text):
        """
        A basic parser to convert markdown sections to Docx.
        In a real scenario, this would be more complex (handling tables, bold, etc.)
        """
        lines = markdown_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headings
            match = re.match(r'^(#+)\s*(.*)', line)
            if match:
                level = len(match.group(1))
                self.add_heading(match.group(2), level=min(level, 9))
                continue
                
            # List items (simple)
            if line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(line[2:], style='List Bullet')
                if self.is_arabic(line):
                    self.set_rtl(p)
                continue

            # Tables (very basic)
            if '|' in line and '-' not in line:
                # This needs a real table builder, but for now we treat as text or skip 
                # as it's hard to parse line-by-line without context.
                # In the final implementation, we might want a proper table parser.
                pass

            self.add_paragraph(line)

    def save(self, output_path):
        self.doc.save(output_path)
        return output_path

def create_report_docx(markdown_content, output_path):
    builder = DocxBuilder()
    builder.build_from_markdown(markdown_content)
    builder.save(output_path)
    return output_path
